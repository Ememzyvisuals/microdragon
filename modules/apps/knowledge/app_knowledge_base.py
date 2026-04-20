"""
microdragon/modules/apps/knowledge/app_knowledge_base.py
═══════════════════════════════════════════════════════════════════════════════
MICRODRAGON APPLICATION KNOWLEDGE BASE
═══════════════════════════════════════════════════════════════════════════════

MICRODRAGON is trained on the complete operational knowledge of every major desktop
application. When MICRODRAGON opens Photoshop, it knows every panel, every shortcut,
every filter, every blending mode, every export option. Same for Excel, Word,
Premiere, Illustrator, and every other app.

This module is the brain behind MICRODRAGON's application control — it tells the
automation engine exactly WHAT to click, WHERE to find it, and WHAT IT DOES.

© 2026 EMEMZYVISUALS DIGITALS — Emmanuel Ariyo
"""

from dataclasses import dataclass, field
from typing import Optional


@dataclass
class AppFeature:
    name: str
    location: str           # Where to find it in the UI
    shortcut: str           # Keyboard shortcut
    description: str        # What it does
    use_case: str           # When to use it


@dataclass
class AppKnowledge:
    app_name: str
    category: str
    version_range: str
    panels: dict            # Panel name → description
    tools: dict             # Tool name → {location, shortcut, description}
    shortcuts: dict         # Action → shortcut key
    workflows: dict         # Common workflow name → steps
    hidden_features: list   # Power user secrets
    automation_notes: str   # How MICRODRAGON specifically controls this app


# ─── Adobe Photoshop ──────────────────────────────────────────────────────────

PHOTOSHOP = AppKnowledge(
    app_name="Adobe Photoshop",
    category="image_editor",
    version_range="CS6 to CC 2025",
    panels={
        "Layers": "Right side — manages all layers, blending modes, opacity",
        "Properties": "Right side — context-sensitive properties for active layer",
        "Adjustments": "Right side — non-destructive adjustments (curves, hue, levels)",
        "History": "Window > History — undo states, create snapshots",
        "Character": "Window > Character — font, size, tracking, kerning",
        "Paragraph": "Window > Paragraph — alignment, indents, spacing",
        "Channels": "Window > Channels — RGB/CMYK channels, alpha channels",
        "Paths": "Window > Paths — vector paths and selections",
        "Swatches": "Window > Swatches — colour swatches",
        "Brushes": "Window > Brushes — brush presets and settings",
        "Tool Options Bar": "Top of screen — changes based on active tool",
        "Options Bar": "Top bar under menu — shows options for selected tool",
        "Info Panel": "Window > Info — pixel values under cursor",
        "Histogram": "Window > Histogram — tonal range of image",
        "Actions": "Window > Actions — recorded macro sequences",
        "Libraries": "Window > Libraries — Creative Cloud shared assets",
    },
    tools={
        "Move Tool": {"shortcut": "V", "location": "Top of toolbar", "desc": "Move layers and selections"},
        "Rectangular Marquee": {"shortcut": "M", "location": "Toolbar", "desc": "Rectangular selection"},
        "Lasso Tool": {"shortcut": "L", "location": "Toolbar", "desc": "Freehand selection"},
        "Magic Wand": {"shortcut": "W", "location": "Toolbar", "desc": "Select by colour similarity"},
        "Quick Selection": {"shortcut": "W", "location": "Toolbar", "desc": "Brush-based auto-selection"},
        "Crop Tool": {"shortcut": "C", "location": "Toolbar", "desc": "Crop and straighten images"},
        "Eyedropper": {"shortcut": "I", "location": "Toolbar", "desc": "Sample colour from image"},
        "Healing Brush": {"shortcut": "J", "location": "Toolbar", "desc": "Remove blemishes, clone nearby texture"},
        "Brush Tool": {"shortcut": "B", "location": "Toolbar", "desc": "Paint with foreground colour"},
        "Clone Stamp": {"shortcut": "S", "location": "Toolbar", "desc": "Copy pixels from one area to another"},
        "History Brush": {"shortcut": "Y", "location": "Toolbar", "desc": "Paint from a history state"},
        "Eraser": {"shortcut": "E", "location": "Toolbar", "desc": "Erase pixels to transparency or background"},
        "Gradient Tool": {"shortcut": "G", "location": "Toolbar", "desc": "Create gradient fills"},
        "Blur/Sharpen/Smudge": {"shortcut": "R", "location": "Toolbar", "desc": "Local focus adjustments"},
        "Dodge/Burn": {"shortcut": "O", "location": "Toolbar", "desc": "Lighten or darken areas"},
        "Pen Tool": {"shortcut": "P", "location": "Toolbar", "desc": "Create vector paths"},
        "Type Tool": {"shortcut": "T", "location": "Toolbar", "desc": "Add text layers"},
        "Path Selection": {"shortcut": "A", "location": "Toolbar", "desc": "Select and move paths"},
        "Rectangle Shape": {"shortcut": "U", "location": "Toolbar", "desc": "Draw vector shapes"},
        "Hand Tool": {"shortcut": "H", "location": "Toolbar", "desc": "Pan around the canvas"},
        "Zoom Tool": {"shortcut": "Z", "location": "Toolbar", "desc": "Zoom in/out on canvas"},
    },
    shortcuts={
        # Essential shortcuts
        "Save": "Ctrl+S / Cmd+S",
        "Save As": "Ctrl+Shift+S",
        "Save for Web": "Ctrl+Alt+Shift+S",
        "Export As": "Ctrl+Alt+Shift+W",
        "Undo": "Ctrl+Z",
        "Step Backward": "Ctrl+Alt+Z",
        "Redo": "Ctrl+Shift+Z",
        "Free Transform": "Ctrl+T",
        "New Layer": "Ctrl+Shift+N",
        "Flatten Image": "Ctrl+Shift+E",
        "Merge Visible": "Ctrl+Shift+E",
        "Duplicate Layer": "Ctrl+J",
        "New Layer via Cut": "Ctrl+Shift+J",
        "Select All": "Ctrl+A",
        "Deselect": "Ctrl+D",
        "Inverse Selection": "Ctrl+Shift+I",
        "Feather Selection": "Shift+F6",
        "Fill with Foreground": "Alt+Backspace",
        "Fill with Background": "Ctrl+Backspace",
        "Open Fill Dialog": "Shift+Backspace",
        "Levels": "Ctrl+L",
        "Curves": "Ctrl+M",
        "Hue/Saturation": "Ctrl+U",
        "Color Balance": "Ctrl+B",
        "Desaturate": "Ctrl+Shift+U",
        "Invert": "Ctrl+I",
        "Canvas Size": "Ctrl+Alt+C",
        "Image Size": "Ctrl+Alt+I",
        "Zoom In": "Ctrl+= / Ctrl++",
        "Zoom Out": "Ctrl+-",
        "Fit to Screen": "Ctrl+0",
        "100% Zoom": "Ctrl+1",
        "New Document": "Ctrl+N",
        "Open Document": "Ctrl+O",
        "Close": "Ctrl+W",
        "Rulers": "Ctrl+R",
        "Guides": "Ctrl+;",
        "Smart Guides": "Ctrl+Alt+;",
        "Layer Group": "Ctrl+G",
        "Ungroup Layers": "Ctrl+Shift+G",
        "Lock Transparent Pixels": "/",
        "Toggle Quick Mask": "Q",
        "Screen Mode": "F",
        "Foreground/Background Colors": "D (reset), X (swap)",
        "Brush Size Increase": "]",
        "Brush Size Decrease": "[",
        "Brush Hardness Increase": "Shift+]",
        "Brush Hardness Decrease": "Shift+[",
        "Increase Opacity 10%": "Number keys (1=10%, 2=20% ... 0=100%)",
        "Smart Object": "Ctrl+Alt+G",
        "Rasterize Layer": "Right-click layer > Rasterize",
        "Clipping Mask": "Ctrl+Alt+G",
        "Layer Mask": "Click mask icon in Layers panel",
        "Refine Edge": "Ctrl+Alt+R (CS6) / Select and Mask (CC)",
        "Object Selection Tool": "W (CC 2020+)",
        "Remove Background": "Right-click layer > Quick Actions > Remove Background",
        "Neural Filters": "Filter > Neural Filters",
        "Content-Aware Fill": "Edit > Content-Aware Fill (after selection)",
        "Liquify": "Filter > Liquify (Shift+Ctrl+X)",
    },
    workflows={
        "Remove background from product photo": [
            "1. Open image (Ctrl+O)",
            "2. Select Object Selection Tool (W)",
            "3. Draw selection around product",
            "4. Click 'Select and Mask' in toolbar",
            "5. Use Edge Detection radius ~3px",
            "6. Output to: New Layer with Layer Mask",
            "7. Delete or hide background layer",
            "8. Export as PNG for transparency"
        ],
        "Create social media post": [
            "1. File > New > choose platform preset (e.g. Instagram Post: 1080x1080)",
            "2. Add background: create gradient or use stock image",
            "3. Add product/subject image on new layer",
            "4. Apply Hue/Saturation adjustment layer for colour grading",
            "5. Add text (T key), use brand fonts",
            "6. Add logo on top layer",
            "7. Export: File > Export > Export As > JPEG quality 85%"
        ],
        "Photo retouching": [
            "1. Duplicate background layer (Ctrl+J) — always work non-destructively",
            "2. Healing Brush (J) for blemishes — Alt+click to sample, paint over blemish",
            "3. Dodge tool (O) to lighten under-eye area, highlights",
            "4. Burn tool (O) to darken, add shadows/contour",
            "5. Curves (Ctrl+M) for global tonal adjustment",
            "6. Hue/Saturation (Ctrl+U) for skin tone",
            "7. Unsharp Mask filter for final sharpening"
        ],
        "Batch resize images": [
            "1. Window > Actions > Create New Action",
            "2. Record: Image > Image Size > set dimensions",
            "3. File > Save for Web (set quality)",
            "4. Stop recording",
            "5. File > Automate > Batch > select Action, source folder, output folder"
        ],
    },
    hidden_features=[
        "Content-Aware Scale (Edit > Content-Aware Scale): stretch canvas without distorting subjects",
        "Vanishing Point (Filter > Vanishing Point): paste/clone in correct perspective",
        "Auto-Align and Auto-Blend Layers for panoramas and focus stacking",
        "Frame Tool (K): create placeholder frames for images",
        "Photomerge: File > Automate > Photomerge for panoramas",
        "Blend If sliders in Layer Style: blend layers based on luminosity without masking",
        "Sky Replacement: Edit > Sky Replacement (CC 2021+)",
        "Pattern Maker: create seamless tiles from any selection",
        "Puppet Warp: Edit > Puppet Warp — place joints and bend/move parts of layer",
        "Gradient Map adjustment: maps tones to colours for cinematic looks",
        "Lab Color mode for non-destructive sharpening (sharpen only the L channel)",
        "Frequency Separation for advanced skin retouching",
        "Camera Raw Filter on any layer: Filter > Camera Raw Filter",
        "Timeline panel for video and GIF animation",
        "3D workspace for text effects and object manipulation",
    ],
    automation_notes="""
MICRODRAGON controls Photoshop via ExtendScript (JSX) for reliable automation:
- Opens Photoshop: subprocess(photoshop_path, '-r', script.jsx)
- All major operations exposed via app.* API
- app.documents.add() creates documents
- app.activeDocument.artLayers operations for layer work
- app.executeAction() for menu commands without UI
- doc.saveAs(file, PNGSaveOptions()) for export
- Batch operations via Batch command or direct loop
"""
)

# ─── Adobe Illustrator ────────────────────────────────────────────────────────

ILLUSTRATOR = AppKnowledge(
    app_name="Adobe Illustrator",
    category="vector_editor",
    version_range="CS6 to CC 2025",
    panels={
        "Layers": "Window > Layers — manage vector layers",
        "Appearance": "Window > Appearance — view and edit object appearance",
        "Swatches": "Window > Swatches — colour management",
        "Brushes": "Window > Brushes — vector brush presets",
        "Symbols": "Window > Symbols — reusable vector symbols",
        "Graphic Styles": "Window > Graphic Styles — saved appearance presets",
        "Artboards": "Window > Artboards — manage multiple artboards",
        "Asset Export": "Window > Asset Export — export individual assets",
        "Character": "Window > Type > Character — typography controls",
        "Paragraph": "Window > Type > Paragraph",
        "OpenType": "Window > Type > OpenType — special characters",
        "Gradient": "Window > Gradient — gradient editor",
        "Color Guide": "Window > Color Guide — harmony rules",
        "Navigator": "Window > Navigator — zoom/pan overview",
        "Info": "Window > Info — coordinates, dimensions",
        "Align": "Window > Align — alignment and distribution",
        "Pathfinder": "Window > Pathfinder — shape boolean operations",
        "Transform": "Window > Transform — precise position/scale",
    },
    tools={
        "Selection Tool": {"shortcut": "V", "location": "Toolbar", "desc": "Select and move objects"},
        "Direct Selection": {"shortcut": "A", "location": "Toolbar", "desc": "Select anchor points"},
        "Magic Wand": {"shortcut": "Y", "location": "Toolbar", "desc": "Select by attribute similarity"},
        "Lasso": {"shortcut": "Q", "location": "Toolbar", "desc": "Freehand select"},
        "Pen Tool": {"shortcut": "P", "location": "Toolbar", "desc": "Create bezier paths"},
        "Add Anchor": {"shortcut": "+", "location": "Toolbar", "desc": "Add anchor to path"},
        "Delete Anchor": {"shortcut": "-", "location": "Toolbar", "desc": "Remove anchor from path"},
        "Anchor Point": {"shortcut": "Shift+C", "location": "Toolbar", "desc": "Convert smooth/corner"},
        "Curvature Tool": {"shortcut": "~", "location": "Toolbar", "desc": "Draw smooth curves easily"},
        "Type Tool": {"shortcut": "T", "location": "Toolbar", "desc": "Add point/area text"},
        "Line Segment": {"shortcut": "\\", "location": "Toolbar", "desc": "Draw straight lines"},
        "Rectangle": {"shortcut": "M", "location": "Toolbar", "desc": "Draw rectangles"},
        "Ellipse": {"shortcut": "L", "location": "Toolbar", "desc": "Draw circles/ellipses"},
        "Polygon": {"shortcut": "None", "location": "Toolbar", "desc": "Draw regular polygons"},
        "Star": {"shortcut": "None", "location": "Toolbar", "desc": "Draw star shapes"},
        "Paintbrush": {"shortcut": "B", "location": "Toolbar", "desc": "Draw with vector brushes"},
        "Pencil": {"shortcut": "N", "location": "Toolbar", "desc": "Freehand path drawing"},
        "Blob Brush": {"shortcut": "Shift+B", "location": "Toolbar", "desc": "Merge painted strokes"},
        "Eraser": {"shortcut": "Shift+E", "location": "Toolbar", "desc": "Erase vector paths"},
        "Scissors": {"shortcut": "C", "location": "Toolbar", "desc": "Cut paths at anchor points"},
        "Knife": {"shortcut": "None", "location": "Toolbar", "desc": "Cut through objects"},
        "Rotate": {"shortcut": "R", "location": "Toolbar", "desc": "Rotate objects"},
        "Reflect": {"shortcut": "O", "location": "Toolbar", "desc": "Mirror/flip objects"},
        "Scale": {"shortcut": "S", "location": "Toolbar", "desc": "Scale objects"},
        "Shear": {"shortcut": "None", "location": "Toolbar", "desc": "Skew objects"},
        "Warp": {"shortcut": "Shift+R", "location": "Toolbar", "desc": "Distort mesh points"},
        "Free Transform": {"shortcut": "E", "location": "Toolbar", "desc": "Transform freely"},
        "Shape Builder": {"shortcut": "Shift+M", "location": "Toolbar", "desc": "Combine/subtract shapes"},
        "Live Paint Bucket": {"shortcut": "K", "location": "Toolbar", "desc": "Fill colour regions"},
        "Gradient": {"shortcut": "G", "location": "Toolbar", "desc": "Apply/edit gradients"},
        "Mesh": {"shortcut": "U", "location": "Toolbar", "desc": "Create gradient mesh"},
        "Eyedropper": {"shortcut": "I", "location": "Toolbar", "desc": "Sample colour/style"},
        "Measure": {"shortcut": "None", "location": "Toolbar", "desc": "Measure distances"},
        "Blend": {"shortcut": "W", "location": "Toolbar", "desc": "Blend between objects"},
        "Symbol Sprayer": {"shortcut": "Shift+S", "location": "Toolbar", "desc": "Spray symbol instances"},
        "Artboard": {"shortcut": "Shift+O", "location": "Toolbar", "desc": "Create/edit artboards"},
        "Slice": {"shortcut": "Shift+K", "location": "Toolbar", "desc": "Define export slices"},
        "Zoom": {"shortcut": "Z", "location": "Toolbar", "desc": "Zoom in/out"},
        "Hand": {"shortcut": "H", "location": "Toolbar", "desc": "Pan canvas"},
    },
    shortcuts={
        "Undo": "Ctrl+Z",
        "Redo": "Ctrl+Shift+Z",
        "Save": "Ctrl+S",
        "Save As": "Ctrl+Shift+S",
        "Export": "Ctrl+Alt+E",
        "New Document": "Ctrl+N",
        "Open": "Ctrl+O",
        "Place (embed image)": "Ctrl+Shift+P",
        "Group": "Ctrl+G",
        "Ungroup": "Ctrl+Shift+G",
        "Lock Object": "Ctrl+2",
        "Lock All Others": "Ctrl+Alt+2",
        "Unlock All": "Ctrl+Alt+2 (again)",
        "Hide Object": "Ctrl+3",
        "Show All": "Ctrl+Alt+3",
        "Bring to Front": "Ctrl+Shift+]",
        "Send to Back": "Ctrl+Shift+[",
        "Bring Forward": "Ctrl+]",
        "Send Backward": "Ctrl+[",
        "Select All": "Ctrl+A",
        "Deselect": "Ctrl+Shift+A",
        "Duplicate in place": "Ctrl+C then Ctrl+F",
        "Paste in Front": "Ctrl+F",
        "Paste in Back": "Ctrl+B",
        "Paste in Place": "Ctrl+Shift+V",
        "Join Paths": "Ctrl+J",
        "Average Points": "Ctrl+Alt+J",
        "Outline Stroke": "Object > Path > Outline Stroke",
        "Expand Appearance": "Object > Expand Appearance",
        "Expand Object": "Object > Expand",
        "Pathfinder Unite": "Alt+click Unite in Pathfinder",
        "Clipping Mask Make": "Ctrl+7",
        "Release Clipping Mask": "Ctrl+Alt+7",
        "Compound Path Make": "Ctrl+8",
        "Type on Path": "Type > Type on a Path > Type on a Path",
        "Create Outlines": "Ctrl+Shift+O",
        "Zoom In": "Ctrl++",
        "Zoom Out": "Ctrl+-",
        "Fit Artboard": "Ctrl+0",
        "Actual Size": "Ctrl+1",
        "Preview Mode": "Ctrl+Y",
        "Align to Artboard": "Use Align panel, click Align to Artboard",
    },
    workflows={
        "Create a logo from scratch": [
            "1. File > New > choose Print or Web preset",
            "2. Use shape tools (M, L, U) for basic geometric shapes",
            "3. Shape Builder (Shift+M) to combine/subtract shapes",
            "4. Add company name: T key, choose font",
            "5. Create Outlines on text: Ctrl+Shift+O",
            "6. Adjust colours in Swatches panel",
            "7. Group all elements: Ctrl+G",
            "8. Export: File > Export > Export As > SVG and PNG"
        ],
        "Create vector icon set": [
            "1. Set up grid artboards (one per icon)",
            "2. Design on pixel grid (View > Pixel Preview)",
            "3. Use Pen tool for complex shapes",
            "4. Shape Builder for unions and subtractions",
            "5. Use consistent stroke width across set",
            "6. Asset Export panel: mark each artboard for export",
            "7. Export all at once: 1x, 2x, 3x SVG and PNG"
        ],
    },
    hidden_features=[
        "Expand Appearance: converts live effects to editable paths",
        "Global Colours: edit one swatch, updates entire document",
        "Variable Width Profile: artistic stroke width variation",
        "Puppet Warp (CC 2021+): distort objects with pin-based warping",
        "Live Corners: hover near rectangle corner → drag to round",
        "Touch Type tool (Shift+T): individual character transformations",
        "Perspective Grid (Shift+P): draw in one/two/three-point perspective",
        "Image Trace: trace raster images to vectors automatically",
        "Recolor Artwork (Edit > Edit Colors): remap all colours at once",
    ],
    automation_notes="""
MICRODRAGON controls Illustrator via ExtendScript (JSX):
- app.documents.add() for new documents
- app.activeDocument.pathItems for vector operations
- Direct manipulation of artboards, layers, objects
- app.executeMenuCommand() for menu access
"""
)

# ─── Adobe Premiere Pro ───────────────────────────────────────────────────────

PREMIERE = AppKnowledge(
    app_name="Adobe Premiere Pro",
    category="video_editor",
    version_range="CC 2019 to CC 2025",
    panels={
        "Project": "Lower left — all imported media assets",
        "Source Monitor": "Upper left — preview raw clips",
        "Program Monitor": "Upper right — preview final edit",
        "Timeline": "Bottom — sequence of clips on tracks",
        "Audio Track Mixer": "Window > Audio Track Mixer — mix audio levels",
        "Lumetri Color": "Window > Lumetri Color — color grading",
        "Lumetri Scopes": "Window > Lumetri Scopes — waveform, vectorscope",
        "Essential Graphics": "Window > Essential Graphics — motion graphics templates",
        "Essential Sound": "Window > Essential Sound — audio repair and mixing",
        "Effects": "Window > Effects — video/audio effects library",
        "Effect Controls": "Window > Effect Controls — parameters of applied effects",
        "History": "Window > History — undo states",
        "Info": "Window > Info — clip information",
        "Markers": "Window > Markers — timeline marker list",
        "Captions and Speech": "Window > Captions and Speech — auto-captions",
    },
    tools={},
    shortcuts={
        "Import": "Ctrl+I",
        "New Sequence": "Ctrl+N",
        "Export Media": "Ctrl+M",
        "Save": "Ctrl+S",
        "Undo": "Ctrl+Z",
        "Play/Pause": "Space",
        "Razor Tool": "C",
        "Selection Tool": "V",
        "Add Edit (razor at playhead)": "Ctrl+K",
        "In Point": "I",
        "Out Point": "O",
        "Clear In": "Ctrl+Shift+I",
        "Clear Out": "Ctrl+Shift+O",
        "Insert into Timeline": ",",
        "Overwrite into Timeline": ".",
        "Ripple Delete": "Shift+Delete",
        "Zoom In Timeline": "=",
        "Zoom Out Timeline": "-",
        "Go to Next Edit Point": "Down Arrow",
        "Go to Previous Edit Point": "Up Arrow",
        "Step Forward 1 Frame": ".",
        "Step Backward 1 Frame": ",",
        "Match Frame": "F",
        "Mark Selection": "/",
        "Lift": ";",
        "Extract": "'",
        "Enable/Disable Clip": "Shift+E",
        "Unlink Audio/Video": "Ctrl+L",
        "Nest Sequence": "Right-click > Nest",
        "Render": "Enter",
        "Audio Gain": "G",
        "Speed/Duration": "Ctrl+R",
    },
    workflows={
        "Basic video edit": [
            "1. File > Import all footage (Ctrl+I)",
            "2. File > New > Sequence (match to footage settings)",
            "3. Drag clips to timeline",
            "4. Razor tool (C) to cut, Selection (V) to move",
            "5. Add transitions: drag from Effects panel",
            "6. Lumetri Color for colour grade",
            "7. Essential Sound for audio cleanup",
            "8. Export: Ctrl+M > H.264 > YouTube preset"
        ],
        "Add auto-captions": [
            "1. Window > Captions and Speech",
            "2. Click 'Transcribe Sequence'",
            "3. Choose language and settings",
            "4. Click 'Create captions' after transcription",
            "5. Edit captions in panel",
            "6. File > Export > Captions for SRT/VTT export"
        ],
    },
    hidden_features=[
        "Auto Reframe: automatically reframe 16:9 to 9:16 for social media",
        "Scene Edit Detection: automatically cut imported video at scene changes",
        "Speech to Text (v22+): free AI captions generation",
        "Proxies: create low-res proxy files for smooth editing of 4K/8K",
        "Track Freeze: freeze frame at current position as still",
        "Adjustment Layers: apply effects to multiple clips at once",
        "Sequence > Auto Match Sequence Settings: one-click sequence optimization",
    ],
    automation_notes="""
MICRODRAGON controls Premiere via ExtendScript (JSX):
- app.project.importFiles() for media import
- app.project.activeSequence for timeline operations
- Track manipulation via sequence.videoTracks / audioTracks
- Export via AMEBatchManager (Adobe Media Encoder)
"""
)

# ─── Adobe After Effects ──────────────────────────────────────────────────────

AFTER_EFFECTS = AppKnowledge(
    app_name="Adobe After Effects",
    category="motion_graphics",
    version_range="CS6 to CC 2025",
    panels={
        "Project": "Left — compositions and footage",
        "Composition": "Centre — preview window",
        "Timeline": "Bottom — layers and keyframes",
        "Effects & Presets": "Right — searchable effects library",
        "Character": "Window > Character — text properties",
        "Paragraph": "Window > Paragraph",
        "Paint": "Window > Paint — brush settings",
        "Brushes": "Window > Brushes",
        "Smoother": "Window > Smoother — smooth keyframe velocity",
        "Motion Sketch": "Window > Motion Sketch — record motion path",
        "Audio": "Window > Audio — audio levels display",
        "Info": "Window > Info — cursor values",
        "Preview": "Window > Preview — preview settings",
    },
    tools={},
    shortcuts={
        "New Composition": "Ctrl+N",
        "Import": "Ctrl+I",
        "New Null Object": "Ctrl+Alt+Shift+Y",
        "New Solid": "Ctrl+Y",
        "New Adjustment Layer": "Ctrl+Alt+Y",
        "New Shape Layer": "No shortcut (use menu)",
        "New Text Layer": "Ctrl+Alt+Shift+T",
        "Duplicate": "Ctrl+D",
        "Pre-compose": "Ctrl+Shift+C",
        "Fit to Comp Width": "Ctrl+Shift+H",
        "Fit to Comp Height": "Ctrl+Shift+G",
        "Solo Layer": "Click S dot in timeline",
        "Lock Layer": "Click lock icon",
        "Collapse Transforms": "Ctrl+E (on precomps)",
        "Reveal Layer Properties": "U (once=keyframes, twice=expressions)",
        "Add Keyframe": "Alt+P (position), Alt+R (rotation), etc.",
        "Set Position Keyframe": "P then click stopwatch",
        "Easy Ease Keyframe": "F9",
        "Set Speed": "Ctrl+Shift+K",
        "Expression Editor": "Alt+click stopwatch",
        "Parent Layer": "Drag parent pick whip to child",
        "Render": "Ctrl+M (add to render queue)",
        "RAM Preview": "0 on numpad (or space bar)",
        "Go to Time": "Ctrl+G",
        "Trim Layer In Point": "[ (at playhead)",
        "Trim Layer Out Point": "] (at playhead)",
        "Increase Timeline Zoom": "=",
        "Decrease Timeline Zoom": "-",
        "Fit Timeline to Window": "Shift+\\",
    },
    workflows={
        "Text animation": [
            "1. Create text layer (Ctrl+Alt+Shift+T)",
            "2. Type your text",
            "3. In timeline, expand text layer > Animator > Add > Position",
            "4. Set range selector Start/End for animate-on effect",
            "5. Add Opacity animator for fade-on",
            "6. Keyframe the Range Selector offset",
            "7. Easy ease all keyframes (F9)"
        ],
        "Logo reveal animation": [
            "1. Import logo PNG/SVG",
            "2. Apply shape as matte or use masks",
            "3. Animate mask path or trim path",
            "4. Add glow effect from Effects & Presets",
            "5. Pre-compose final logo with background",
            "6. Add camera with depth of field for 3D feel"
        ],
    },
    hidden_features=[
        "Expressions: link properties with JavaScript-like expressions",
        "Motion Bro: free plugin for After Effects templates",
        "Trim Paths: animate stroke drawing on shape layers",
        "Puppet Pin Tool: distort images/characters with joints",
        "Content-Aware Fill (CC 2019+): remove objects from video",
        "Essential Graphics panel: create editable MOGRT templates for Premiere",
        "3D Camera Tracker: track 3D space from 2D footage",
        "Roto Brush 2 (CC 2020+): AI-powered rotoscoping",
    ],
    automation_notes="""
MICRODRAGON controls After Effects via ExtendScript (JSX):
- app.project.importFile() for media
- app.project.items for composition operations
- Layer property manipulation via layer.property()
- Keyframe creation via property.addKey()
- Render queue via app.project.renderQueue
"""
)

# ─── Microsoft Excel ──────────────────────────────────────────────────────────

EXCEL = AppKnowledge(
    app_name="Microsoft Excel",
    category="spreadsheet",
    version_range="2016 to 365 (2025)",
    panels={
        "Ribbon": "Top — all commands organised in tabs",
        "Formula Bar": "Below ribbon — shows/edit cell formula",
        "Name Box": "Left of formula bar — shows cell address",
        "Sheet Tabs": "Bottom — navigate between worksheets",
        "Status Bar": "Bottom — SUM, AVERAGE, COUNT of selection",
        "Mini Toolbar": "Appears on right-click — quick format options",
    },
    tools={},
    shortcuts={
        "New Workbook": "Ctrl+N",
        "Open": "Ctrl+O",
        "Save": "Ctrl+S",
        "Save As": "F12",
        "Close": "Ctrl+W",
        "Print": "Ctrl+P",
        "Undo": "Ctrl+Z",
        "Redo": "Ctrl+Y",
        "Find": "Ctrl+F",
        "Replace": "Ctrl+H",
        "Go To": "Ctrl+G or F5",
        "Select All": "Ctrl+A",
        "Select Column": "Ctrl+Space",
        "Select Row": "Shift+Space",
        "Copy": "Ctrl+C",
        "Cut": "Ctrl+X",
        "Paste": "Ctrl+V",
        "Paste Special": "Ctrl+Alt+V",
        "Insert Row": "Ctrl+Shift++",
        "Delete Row": "Ctrl+-",
        "Bold": "Ctrl+B",
        "Italic": "Ctrl+I",
        "Underline": "Ctrl+U",
        "Format Cells": "Ctrl+1",
        "Number Format": "Ctrl+Shift+1",
        "Date Format": "Ctrl+Shift+3",
        "Currency Format": "Ctrl+Shift+4",
        "Percentage Format": "Ctrl+Shift+5",
        "AutoSum": "Alt+=",
        "Enter formula": "= then type formula",
        "Complete entry move down": "Enter",
        "Complete entry move right": "Tab",
        "Complete entry stay": "Ctrl+Enter",
        "Start new line in cell": "Alt+Enter",
        "Fill Down": "Ctrl+D",
        "Fill Right": "Ctrl+R",
        "Toggle Absolute Reference": "F4 (cycles $A$1, A$1, $A1, A1)",
        "Calculate Now": "F9",
        "Insert Function": "Shift+F3",
        "Create Named Range": "Ctrl+Shift+F3 or Name Box",
        "Show All Formulas": "Ctrl+`",
        "Expand/Collapse groups": "Alt+Shift+= / Alt+Shift+-",
        "Group Rows/Cols": "Alt+Shift+Right",
        "Ungroup": "Alt+Shift+Left",
        "Create Table": "Ctrl+T",
        "Insert Pivot Table": "Alt+N+V",
        "Freeze Panes": "Alt+W+F+F",
        "Filter": "Ctrl+Shift+L",
        "Flash Fill": "Ctrl+E",
        "New Sheet": "Shift+F11",
        "Move to last used cell": "Ctrl+End",
        "Move to A1": "Ctrl+Home",
        "Select to last cell in column": "Ctrl+Shift+Down",
        "Select to last cell in row": "Ctrl+Shift+Right",
        "Zoom to Selection": "Alt+W+G",
    },
    workflows={
        "Create a monthly budget tracker": [
            "1. Create header row: Category, Budget, Actual, Variance",
            "2. List expense categories in column A",
            "3. Enter budget amounts in column B",
            "4. Actual amounts in column C (updated monthly)",
            "5. Variance = =B2-C2 in column D",
            "6. Add SUM row at bottom: =SUM(B2:B20) etc.",
            "7. Conditional formatting: red for negative variance",
            "8. Create bar chart: select data > Insert > Bar Chart"
        ],
        "VLOOKUP to match data": [
            "1. Ensure lookup table has key in first column",
            "2. In result cell: =VLOOKUP(lookup_value, table_array, col_index, FALSE)",
            "3. FALSE = exact match (recommended for IDs/codes)",
            "4. Or use XLOOKUP (365): =XLOOKUP(value, lookup_range, return_range)"
        ],
        "Pivot table from sales data": [
            "1. Click any cell in data range",
            "2. Insert > PivotTable > New Worksheet",
            "3. Drag 'Date' to Rows (right-click > Group by Month)",
            "4. Drag 'Revenue' to Values (Sum)",
            "5. Drag 'Region' to Columns",
            "6. Insert > PivotChart for visualisation"
        ],
    },
    hidden_features=[
        "XLOOKUP: replaces VLOOKUP with more power (365/2021+)",
        "LET function: define variables in formulas",
        "LAMBDA function: create custom reusable functions without VBA",
        "Dynamic Arrays: SORT, FILTER, UNIQUE, SEQUENCE, RANDARRAY",
        "Flash Fill (Ctrl+E): auto-detect patterns and fill column",
        "Power Query: Get & Transform Data — import, clean, reshape data",
        "Power Pivot: create data models across multiple tables",
        "Solver: What-if analysis to find optimal solution",
        "Forecast Sheet: Insert > Forecast Sheet for trend projections",
        "Ideas/Analyze Data: Home > Analyze Data — AI suggests charts and insights",
        "Watch Window: monitor cells from other sheets",
        "Name Manager: Ctrl+F3 — manage all named ranges",
        "INDIRECT function: create dynamic cell references from text strings",
        "Data Validation: create dropdowns and input rules",
        "Consolidate: Data > Consolidate — combine ranges from multiple sheets",
    ],
    automation_notes="""
MICRODRAGON controls Excel via:
1. openpyxl (Python) — read/write .xlsx without Excel installed
2. VBA macros via PowerShell/COM on Windows
3. xlwings for Python-Excel bridge
Excel VBA automation:
  Set xl = CreateObject("Excel.Application")
  Set wb = xl.Workbooks.Open(path)
  wb.ActiveSheet.Cells(1,1).Value = "MICRODRAGON"
  wb.Save
"""
)

# ─── Microsoft Word ───────────────────────────────────────────────────────────

WORD = AppKnowledge(
    app_name="Microsoft Word",
    category="word_processor",
    version_range="2016 to 365 (2025)",
    panels={
        "Ribbon": "All commands in tabbed interface",
        "Ruler": "View > Ruler — shows indent and tab markers",
        "Navigation Pane": "View > Navigation Pane — document headings",
        "Styles Gallery": "Home tab — quick styles application",
        "Review Pane": "Review > Reviewing Pane — all tracked changes",
        "Document Map": "Navigation Pane in heading view",
    },
    tools={},
    shortcuts={
        "Bold": "Ctrl+B",
        "Italic": "Ctrl+I",
        "Underline": "Ctrl+U",
        "Strikethrough": "Alt+H, 4",
        "Subscript": "Ctrl+=",
        "Superscript": "Ctrl+Shift+=",
        "Increase Font Size": "Ctrl+Shift+>",
        "Decrease Font Size": "Ctrl+Shift+<",
        "Align Left": "Ctrl+L",
        "Centre": "Ctrl+E",
        "Align Right": "Ctrl+R",
        "Justify": "Ctrl+J",
        "Heading 1": "Ctrl+Alt+1",
        "Heading 2": "Ctrl+Alt+2",
        "Heading 3": "Ctrl+Alt+3",
        "Normal Style": "Ctrl+Shift+N",
        "Bullet List": "Ctrl+Shift+L",
        "Insert Link": "Ctrl+K",
        "Find": "Ctrl+F",
        "Replace": "Ctrl+H",
        "Go To": "Ctrl+G",
        "Select All": "Ctrl+A",
        "Copy Formatting": "Ctrl+Shift+C",
        "Paste Formatting": "Ctrl+Shift+V",
        "Track Changes": "Ctrl+Shift+E",
        "Accept All Changes": "Review > Accept > Accept All",
        "Insert Page Break": "Ctrl+Enter",
        "Insert Section Break": "Layout > Breaks",
        "Table of Contents": "References > Table of Contents",
        "Spelling & Grammar": "F7",
        "Word Count": "Review > Word Count",
        "Save as PDF": "File > Export > Create PDF/XPS",
        "Print Preview": "Ctrl+F2",
        "Zoom": "View > Zoom",
        "Show/Hide Marks": "Ctrl+Shift+8",
        "Insert Comment": "Ctrl+Alt+M",
        "Insert Footnote": "Ctrl+Alt+F",
        "Insert Endnote": "Ctrl+Alt+D",
    },
    workflows={},
    hidden_features=[
        "Compare Documents: Review > Compare — shows differences between two versions",
        "Document Protection: Review > Protect Document",
        "Mail Merge: Mailings tab — create personalised bulk documents",
        "Building Blocks: Insert > Quick Parts — reuse text blocks",
        "AutoCorrect: File > Options > Proofing > AutoCorrect Options — custom shortcuts",
        "Styles and Themes: completely controls document appearance",
        "Cross-references: Insert > Cross-reference — links to figures, headings, page numbers",
        "Accessibility Checker: Review > Check Accessibility",
        "Resume Assistant (365): Review > Resume Assistant — LinkedIn integration",
        "Editor (365): Home > Editor — advanced grammar and style checking",
        "Focus Mode: View > Focus — distraction-free writing",
        "Dark Mode: supported in 365",
    ],
    automation_notes="""
MICRODRAGON controls Word via python-docx or COM automation:
python-docx:
  from docx import Document
  doc = Document()
  doc.add_heading('Title', level=0)
  doc.add_paragraph('Content')
  doc.save('output.docx')
COM (Windows):
  word = win32.Dispatch('Word.Application')
  doc = word.Documents.Open(path)
"""
)

# ─── Registry ─────────────────────────────────────────────────────────────────

APP_KNOWLEDGE_BASE = {
    "photoshop": PHOTOSHOP,
    "adobe photoshop": PHOTOSHOP,
    "illustrator": ILLUSTRATOR,
    "adobe illustrator": ILLUSTRATOR,
    "premiere": PREMIERE,
    "premiere pro": PREMIERE,
    "adobe premiere": PREMIERE,
    "after effects": AFTER_EFFECTS,
    "excel": EXCEL,
    "microsoft excel": EXCEL,
    "word": WORD,
    "microsoft word": WORD,
}


def get_app_knowledge(app_name: str) -> AppKnowledge | None:
    """Get knowledge base entry for an application."""
    return APP_KNOWLEDGE_BASE.get(app_name.lower())


def get_shortcut(app_name: str, action: str) -> str:
    """Get keyboard shortcut for an action in an application."""
    app = get_app_knowledge(app_name)
    if not app:
        return f"App '{app_name}' not in knowledge base"
    shortcut = app.shortcuts.get(action)
    if not shortcut:
        # Fuzzy search
        matches = [(k, v) for k, v in app.shortcuts.items()
                   if action.lower() in k.lower()]
        if matches:
            return "\n".join(f"{k}: {v}" for k, v in matches[:3])
        return f"No shortcut found for '{action}' in {app_name}"
    return f"{action}: {shortcut}"


def get_workflow(app_name: str, workflow_name: str) -> list:
    """Get step-by-step workflow for a task in an application."""
    app = get_app_knowledge(app_name)
    if not app:
        return [f"App '{app_name}' not in knowledge base"]
    for name, steps in app.workflows.items():
        if workflow_name.lower() in name.lower():
            return steps
    return [f"Workflow '{workflow_name}' not found for {app_name}"]


def build_app_context_prompt(app_name: str, task: str) -> str:
    """Build a rich context prompt for AI when operating an application."""
    app = get_app_knowledge(app_name)
    if not app:
        return f"Task: {task} in {app_name}"

    relevant_shortcuts = {k: v for k, v in list(app.shortcuts.items())[:20]}
    relevant_panels = list(app.panels.items())[:8]

    return f"""You are controlling {app.app_name} to complete this task: {task}

APPLICATION KNOWLEDGE:
- Version: {app.version_range}
- Category: {app.category}

KEY PANELS:
{chr(10).join(f'  {name}: {desc}' for name, desc in relevant_panels)}

KEY SHORTCUTS:
{chr(10).join(f'  {k}: {v}' for k, v in relevant_shortcuts.items())}

AUTOMATION APPROACH:
{app.automation_notes}

Complete the task step by step. Use the shortest path.
Reference specific panel names, tool names, and shortcuts.
"""
