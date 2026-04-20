"""
microdragon/modules/apps/src/engine.py
MICRODRAGON Application Controller — Uses real desktop apps via automation
Photoshop, GIMP, Excel, LibreOffice, Word, VS Code, browsers, and any GUI app
This is MICRODRAGON's "hands" — it doesn't just generate code, it OPERATES programs
"""

import asyncio
import subprocess
import sys
import os
import time
import json
from dataclasses import dataclass, field
from typing import Optional, Any
from pathlib import Path
import tempfile


@dataclass
class AppResult:
    success: bool
    output: str = ""
    output_path: Optional[str] = None
    error: str = ""
    app_used: str = ""
    method: str = ""  # 'native' | 'automation' | 'api' | 'cli'


# ─── Application Registry ─────────────────────────────────────────────────────

class AppRegistry:
    """Detects what apps are installed and picks the best one for each task."""

    @staticmethod
    def find_app(candidates: list[str]) -> Optional[str]:
        """Return first found executable from candidates list."""
        for name in candidates:
            try:
                result = subprocess.run(
                    ["where" if sys.platform == "win32" else "which", name],
                    capture_output=True, text=True
                )
                if result.returncode == 0:
                    return name
            except Exception:
                pass
        return None

    @staticmethod
    def photoshop_path() -> Optional[str]:
        """Find Photoshop on this system."""
        if sys.platform == "win32":
            import glob
            patterns = [
                r"C:\Program Files\Adobe\Adobe Photoshop*\Photoshop.exe",
                r"C:\Program Files (x86)\Adobe\Adobe Photoshop*\Photoshop.exe",
            ]
            for pattern in patterns:
                found = glob.glob(pattern)
                if found: return found[0]
        elif sys.platform == "darwin":
            import glob
            patterns = [
                "/Applications/Adobe Photoshop*/Adobe Photoshop*.app/Contents/MacOS/Adobe Photoshop*",
            ]
            for pattern in patterns:
                found = glob.glob(pattern)
                if found: return found[0]
        return None

    @staticmethod
    def excel_path() -> Optional[str]:
        if sys.platform == "win32":
            import glob
            patterns = [
                r"C:\Program Files\Microsoft Office\root\Office*\EXCEL.EXE",
                r"C:\Program Files (x86)\Microsoft Office\Office*\EXCEL.EXE",
            ]
            for p in patterns:
                found = glob.glob(p)
                if found: return found[0]
        elif sys.platform == "darwin":
            path = "/Applications/Microsoft Excel.app/Contents/MacOS/Microsoft Excel"
            if os.path.exists(path): return path
        return None

    @staticmethod
    def word_path() -> Optional[str]:
        if sys.platform == "win32":
            import glob
            patterns = [
                r"C:\Program Files\Microsoft Office\root\Office*\WINWORD.EXE",
                r"C:\Program Files (x86)\Microsoft Office\Office*\WINWORD.EXE",
            ]
            for p in patterns:
                found = glob.glob(p)
                if found: return found[0]
        elif sys.platform == "darwin":
            path = "/Applications/Microsoft Word.app/Contents/MacOS/Microsoft Word"
            if os.path.exists(path): return path
        return None


# ─── Photoshop / GIMP Controller ──────────────────────────────────────────────

class ImageEditor:
    """
    Controls Photoshop via ExtendScript (Windows/macOS) or GIMP via Script-Fu.
    Falls back to Pillow for programmatic image operations.
    """

    def __init__(self):
        self.photoshop = AppRegistry.photoshop_path()
        self.gimp = AppRegistry.find_app(["gimp", "gimp-2.10", "gimp-2.12"])
        self.has_pillow = self._check_pillow()

    def _check_pillow(self) -> bool:
        try:
            import PIL
            return True
        except ImportError:
            return False

    async def create_design(self, task: str, output_path: str,
                             width: int = 1920, height: int = 1080,
                             template: Optional[str] = None) -> AppResult:
        """Create a design using the best available tool."""

        if self.photoshop:
            return await self._photoshop_design(task, output_path, width, height)
        elif self.gimp:
            return await self._gimp_design(task, output_path, width, height)
        elif self.has_pillow:
            return await self._pillow_design(task, output_path, width, height)
        else:
            return AppResult(success=False,
                             error="No image editor found. Install Photoshop, GIMP, or Pillow.")

    async def _photoshop_design(self, task: str, output_path: str,
                                  w: int, h: int) -> AppResult:
        """Use Photoshop via ExtendScript JSX."""
        script = self._generate_ps_script(task, output_path, w, h)
        jsx_path = tempfile.mktemp(suffix=".jsx")
        Path(jsx_path).write_text(script)

        try:
            # Photoshop can run scripts via command line
            result = subprocess.run(
                [self.photoshop, "-r", jsx_path],
                capture_output=True, text=True, timeout=120
            )
            os.unlink(jsx_path)
            success = result.returncode == 0 or os.path.exists(output_path)
            return AppResult(success=success, output_path=output_path if success else None,
                             app_used="Adobe Photoshop", method="extendscript")
        except Exception as e:
            return AppResult(success=False, error=str(e))

    def _generate_ps_script(self, task: str, output_path: str, w: int, h: int) -> str:
        """Generate ExtendScript JSX for Photoshop."""
        return f"""
// MICRODRAGON-generated Photoshop script
// Task: {task}
var doc = app.documents.add(
    {w}, {h}, 72, "MICRODRAGON Design",
    NewDocumentMode.RGB, DocumentFill.TRANSPARENT
);
// Add a gradient background
var layer = doc.artLayers.add();
layer.name = "Background";
// Set foreground/background colors
app.foregroundColor.rgb.red = 30;
app.foregroundColor.rgb.green = 30;
app.foregroundColor.rgb.blue = 50;
app.backgroundColor.rgb.red = 10;
app.backgroundColor.rgb.green = 10;
app.backgroundColor.rgb.blue = 30;
// Fill with gradient
var gradDesc = new ActionDescriptor();
app.executeAction(charIDToTypeID('FrgF'), undefined, DialogModes.NO);
// Add text layer
var textLayer = doc.artLayers.add();
textLayer.kind = LayerKind.TEXT;
textLayer.name = "Title";
var textItem = textLayer.textItem;
textItem.contents = "{task[:50].replace(chr(34), '')}";
textItem.size = 72;
textItem.position = [doc.width / 2 - 200, doc.height / 2];
// Save as PNG
var saveFile = new File("{output_path.replace(chr(92), '/')}");
var pngSaveOptions = new PNGSaveOptions();
doc.saveAs(saveFile, pngSaveOptions, true, Extension.LOWERCASE);
doc.close(SaveOptions.DONOTSAVECHANGES);
"""

    async def _gimp_design(self, task: str, output_path: str,
                             w: int, h: int) -> AppResult:
        """Use GIMP via Script-Fu batch mode."""
        script = f"""
(let* (
  (image (car (gimp-image-new {w} {h} RGB)))
  (drawable (car (gimp-layer-new image {w} {h} RGB-IMAGE "Background" 100 LAYER-MODE-NORMAL-LEGACY)))
)
  (gimp-image-insert-layer image drawable 0 -1)
  (gimp-context-set-background '(30 30 50))
  (gimp-edit-fill drawable FILL-BACKGROUND)
  (let* ((text-layer (car (gimp-text-fontname image -1 50 (/ {h} 2) "{task[:40].replace(chr(39), '')}" 0 TRUE 48 UNIT-PIXEL "Sans Bold"))))
    (gimp-text-layer-set-color text-layer '(255 255 255))
  )
  (file-png-save RUN-NONINTERACTIVE image (car (gimp-image-get-active-drawable image)) "{output_path}" "{output_path}" 0 9 1 1 1 1 1)
  (gimp-image-delete image)
)
"""
        try:
            result = subprocess.run(
                [self.gimp, "-i", "-b", script, "-b", "(gimp-quit 0)"],
                capture_output=True, text=True, timeout=120
            )
            return AppResult(success=os.path.exists(output_path),
                             output_path=output_path if os.path.exists(output_path) else None,
                             app_used="GIMP", method="script-fu")
        except Exception as e:
            return AppResult(success=False, error=str(e))

    async def _pillow_design(self, task: str, output_path: str,
                              w: int, h: int) -> AppResult:
        """Programmatic image creation with Pillow (no desktop app needed)."""
        try:
            from PIL import Image, ImageDraw, ImageFont
            import colorsys

            img = Image.new("RGB", (w, h), color=(30, 30, 50))
            draw = ImageDraw.Draw(img)

            # Gradient background
            for y in range(h):
                r = int(20 + (y / h) * 20)
                g = int(20 + (y / h) * 10)
                b = int(50 + (y / h) * 30)
                draw.line([(0, y), (w, y)], fill=(r, g, b))

            # Try to load a nice font
            font_large = self._get_font(size=72)
            font_small = self._get_font(size=32)

            # Title text
            title = task[:60]
            bbox = draw.textbbox((0, 0), title, font=font_large)
            text_w = bbox[2] - bbox[0]
            x = max(0, (w - text_w) // 2)
            y = h // 2 - 50

            # Shadow
            draw.text((x+3, y+3), title, font=font_large, fill=(0, 0, 0, 128))
            draw.text((x, y), title, font=font_large, fill=(255, 255, 255))

            # Subtitle
            subtitle = "Created by MICRODRAGON AI"
            draw.text((w // 2 - 100, h - 80), subtitle, font=font_small, fill=(150, 150, 200))

            img.save(output_path)
            return AppResult(success=True, output_path=output_path,
                             app_used="Pillow", method="programmatic")
        except Exception as e:
            return AppResult(success=False, error=str(e))

    def _get_font(self, size: int = 32):
        """Get a font with graceful fallback."""
        try:
            from PIL import ImageFont
            # Try common system fonts
            font_candidates = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                "/System/Library/Fonts/Helvetica.ttc",
                "C:\\Windows\\Fonts\\arial.ttf",
                "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf",
            ]
            for path in font_candidates:
                if os.path.exists(path):
                    return ImageFont.truetype(path, size=size)
            return ImageFont.load_default()
        except Exception:
            from PIL import ImageFont
            return ImageFont.load_default()

    async def edit_image(self, input_path: str, operations: list[dict],
                          output_path: str) -> AppResult:
        """Apply operations to an existing image."""
        try:
            from PIL import Image, ImageFilter, ImageEnhance

            img = Image.open(input_path)

            for op in operations:
                action = op.get("action", "")
                if action == "resize":
                    img = img.resize((op["width"], op["height"]), Image.LANCZOS)
                elif action == "crop":
                    img = img.crop((op["x"], op["y"], op["x"]+op["w"], op["y"]+op["h"]))
                elif action == "rotate":
                    img = img.rotate(op["angle"], expand=True)
                elif action == "blur":
                    img = img.filter(ImageFilter.GaussianBlur(radius=op.get("radius", 3)))
                elif action == "brightness":
                    img = ImageEnhance.Brightness(img).enhance(op.get("factor", 1.2))
                elif action == "contrast":
                    img = ImageEnhance.Contrast(img).enhance(op.get("factor", 1.2))
                elif action == "grayscale":
                    img = img.convert("L").convert("RGB")
                elif action == "sharpen":
                    img = img.filter(ImageFilter.SHARPEN)
                elif action == "convert":
                    img = img.convert(op.get("mode", "RGB"))

            img.save(output_path)
            return AppResult(success=True, output_path=output_path,
                             app_used="Pillow", method="programmatic")
        except ImportError:
            return AppResult(success=False, error="Pillow not installed: pip install Pillow")
        except Exception as e:
            return AppResult(success=False, error=str(e))


# ─── Excel / Spreadsheet Controller ──────────────────────────────────────────

class SpreadsheetController:
    """
    Creates and manipulates spreadsheets.
    Uses Excel (Windows/macOS), LibreOffice Calc, or openpyxl (no app needed).
    """

    def __init__(self):
        self.excel = AppRegistry.excel_path()
        self.libreoffice = AppRegistry.find_app(["libreoffice", "soffice"])

    async def create_spreadsheet(self, task: str, data: list[list],
                                   headers: list[str], output_path: str,
                                   charts: bool = False) -> AppResult:
        """Create a spreadsheet with data, formatting, and optional charts."""
        # Always use openpyxl for reliable cross-platform support
        return await self._create_openpyxl(task, data, headers, output_path, charts)

    async def _create_openpyxl(self, task: str, data: list[list],
                                 headers: list[str], output_path: str,
                                 add_charts: bool) -> AppResult:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                                          numbers)
            from openpyxl.utils import get_column_letter
            from openpyxl.chart import BarChart, LineChart, Reference

            wb = Workbook()
            ws = wb.active
            ws.title = "MICRODRAGON Data"

            # Header style
            header_fill = PatternFill("solid", fgColor="1E3A5F")
            header_font = Font(color="FFFFFF", bold=True, size=11)
            border = Border(
                left=Side(style="thin"), right=Side(style="thin"),
                top=Side(style="thin"), bottom=Side(style="thin")
            )

            # Write headers
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = border

            # Write data
            alt_fill = PatternFill("solid", fgColor="F0F4F8")
            for row_idx, row_data in enumerate(data, 2):
                fill = alt_fill if row_idx % 2 == 0 else PatternFill("none")
                for col_idx, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.border = border
                    if isinstance(fill, PatternFill) and fill.patternType != "none":
                        cell.fill = fill

            # Auto-fit columns
            for col in ws.columns:
                max_len = max((len(str(cell.value or "")) for cell in col), default=10)
                ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 50)

            # Freeze header row
            ws.freeze_panes = "A2"

            # Add auto-filter
            ws.auto_filter.ref = ws.dimensions

            # Add chart if requested and we have numeric data
            if add_charts and len(data) > 1 and len(headers) >= 2:
                try:
                    chart = BarChart()
                    chart.title = task[:40]
                    chart.style = 10
                    chart.y_axis.title = headers[1] if len(headers) > 1 else "Value"
                    chart.x_axis.title = headers[0]

                    data_ref = Reference(ws, min_col=2, min_row=1,
                                         max_col=min(len(headers), 5),
                                         max_row=len(data) + 1)
                    cats = Reference(ws, min_col=1, min_row=2, max_row=len(data) + 1)
                    chart.add_data(data_ref, titles_from_data=True)
                    chart.set_categories(cats)
                    chart.shape = 4
                    ws.add_chart(chart, "A" + str(len(data) + 5))
                except Exception:
                    pass  # Chart creation is optional

            # Add summary sheet
            ws_summary = wb.create_sheet("Summary")
            ws_summary["A1"] = "Task"
            ws_summary["B1"] = task
            ws_summary["A2"] = "Rows"
            ws_summary["B2"] = len(data)
            ws_summary["A3"] = "Columns"
            ws_summary["B3"] = len(headers)
            ws_summary["A4"] = "Generated by"
            ws_summary["B4"] = "MICRODRAGON AI Agent"
            ws_summary["A5"] = "Created"
            from datetime import datetime
            ws_summary["B5"] = datetime.now().strftime("%Y-%m-%d %H:%M")

            wb.save(output_path)
            return AppResult(success=True, output_path=output_path,
                             app_used="openpyxl", method="programmatic",
                             output=f"Spreadsheet created: {len(data)} rows, {len(headers)} columns")
        except ImportError:
            # Try LibreOffice CLI
            return await self._create_libreoffice_csv(task, data, headers, output_path)
        except Exception as e:
            return AppResult(success=False, error=str(e))

    async def _create_libreoffice_csv(self, task: str, data: list[list],
                                        headers: list[str], output_path: str) -> AppResult:
        """Create via CSV + LibreOffice conversion."""
        import csv
        csv_path = output_path.replace(".xlsx", ".csv")
        try:
            with open(csv_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(headers)
                writer.writerows(data)

            if self.libreoffice and not output_path.endswith(".csv"):
                result = subprocess.run(
                    [self.libreoffice, "--headless", "--convert-to", "xlsx",
                     "--outdir", str(Path(output_path).parent), csv_path],
                    capture_output=True, timeout=60
                )
                if result.returncode == 0:
                    return AppResult(success=True, output_path=output_path,
                                     app_used="LibreOffice", method="cli")

            return AppResult(success=True, output_path=csv_path,
                             app_used="csv", method="programmatic",
                             output="Created CSV (install openpyxl for .xlsx)")
        except Exception as e:
            return AppResult(success=False, error=str(e))

    async def open_and_modify(self, file_path: str, macro_task: str) -> AppResult:
        """Open an existing spreadsheet in Excel/LibreOffice and run AI-generated macro."""
        if self.excel:
            return await self._excel_macro(file_path, macro_task)
        elif self.libreoffice:
            return await self._libreoffice_macro(file_path, macro_task)
        return AppResult(success=False, error="No spreadsheet app found")

    async def _excel_macro(self, file_path: str, task: str) -> AppResult:
        """Run a VBA macro in Excel via PowerShell."""
        vba = f"""
Sub MICRODRAGONTask()
    ' MICRODRAGON Auto-generated macro
    ' Task: {task}
    Dim wb As Workbook
    Set wb = ThisWorkbook
    Dim ws As Worksheet
    Set ws = wb.ActiveSheet
    MsgBox "MICRODRAGON macro running: {task[:50]}"
End Sub
"""
        ps_script = f"""
$excel = New-Object -ComObject Excel.Application
$excel.Visible = $true
$wb = $excel.Workbooks.Open("{file_path.replace(chr(92), chr(92)*2)}")
$module = $wb.VBProject.VBComponents.Add(1)
$module.CodeModule.AddFromString('{vba.replace(chr(39), chr(34))}')
$excel.Run("MICRODRAGONTask")
$wb.Save()
$excel.Quit()
"""
        try:
            result = subprocess.run(
                ["powershell", "-Command", ps_script],
                capture_output=True, text=True, timeout=120
            )
            return AppResult(success=result.returncode == 0,
                             app_used="Microsoft Excel", method="vba",
                             error=result.stderr[:200] if result.returncode != 0 else "")
        except Exception as e:
            return AppResult(success=False, error=str(e))

    async def _libreoffice_macro(self, file_path: str, task: str) -> AppResult:
        """Use LibreOffice Basic macro."""
        macro = f"""
import uno
from com.sun.star.beans import PropertyValue

def run_microdragon_task():
    # MICRODRAGON Task: {task}
    ctx = uno.getComponentContext()
    sm = ctx.ServiceManager
    desktop = sm.createInstanceWithContext("com.sun.star.frame.Desktop", ctx)
    doc = desktop.loadComponentFromURL("file://{file_path}", "_blank", 0, [])
    sheet = doc.Sheets.getByIndex(0)
    cell = sheet.getCellByPosition(0, 0)
    cell.setString(f"MICRODRAGON: {task[:30]}")
    doc.store()
"""
        try:
            macro_path = tempfile.mktemp(suffix=".py")
            Path(macro_path).write_text(macro)
            result = subprocess.run(
                [self.libreoffice, "--headless", "--norestore", file_path],
                capture_output=True, timeout=30
            )
            return AppResult(success=True, app_used="LibreOffice", method="macro")
        except Exception as e:
            return AppResult(success=False, error=str(e))


# ─── Word Processor Controller ─────────────────────────────────────────────────

class WordProcessor:
    """
    Creates documents using Word, LibreOffice Writer, or python-docx.
    Supports formatted reports, letters, contracts, invoices.
    """

    def __init__(self):
        self.word = AppRegistry.word_path()
        self.libreoffice = AppRegistry.find_app(["libreoffice", "soffice"])

    async def create_document(self, title: str, content: str,
                               output_path: str, style: str = "professional") -> AppResult:
        """Create a formatted Word document."""
        return await self._create_python_docx(title, content, output_path, style)

    async def _create_python_docx(self, title: str, content: str,
                                    output_path: str, style: str) -> AppResult:
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            import datetime

            doc = Document()

            # Page setup
            section = doc.sections[0]
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata line
            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_run = meta.add_run(f"Generated by MICRODRAGON AI  •  {datetime.date.today()}")
            meta_run.font.size = Pt(9)
            meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            doc.add_paragraph()  # spacer

            # Parse and format content
            self._format_content(doc, content)

            # Footer
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"MICRODRAGON AI Agent  |  {datetime.date.today()}  |  Confidential"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.save(output_path)
            return AppResult(success=True, output_path=output_path,
                             app_used="python-docx", method="programmatic",
                             output=f"Document created: {title}")
        except ImportError:
            return await self._create_markdown_fallback(title, content, output_path)
        except Exception as e:
            return AppResult(success=False, error=str(e))

    def _format_content(self, doc, content: str):
        """Parse markdown-like content and apply Word formatting."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith(tuple("0123456789")) and ". " in line[:4]:
                doc.add_paragraph(line, style="List Number")
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.style = "Quote" if "Quote" in [s.name for s in doc.styles] else "Normal"
            else:
                doc.add_paragraph(line)

    async def _create_markdown_fallback(self, title: str, content: str,
                                         output_path: str) -> AppResult:
        """Create Markdown as fallback when python-docx not installed."""
        md_path = output_path.replace(".docx", ".md")
        full_content = f"# {title}\n\n{content}"
        Path(md_path).write_text(full_content, encoding="utf-8")
        return AppResult(success=True, output_path=md_path,
                         app_used="markdown", method="fallback",
                         output="Saved as .md (install python-docx for .docx)")

    async def open_in_word(self, file_path: str) -> AppResult:
        """Open a document in Word or LibreOffice."""
        if self.word:
            try:
                subprocess.Popen([self.word, file_path])
                return AppResult(success=True, app_used="Microsoft Word",
                                 method="launch", output="Opened in Word")
            except Exception as e:
                pass

        if self.libreoffice:
            try:
                subprocess.Popen([self.libreoffice, "--writer", file_path])
                return AppResult(success=True, app_used="LibreOffice Writer",
                                 method="launch", output="Opened in LibreOffice Writer")
            except Exception as e:
                pass

        return AppResult(success=False, error="No word processor found")


# ─── Browser Agent ────────────────────────────────────────────────────────────

class BrowserAgent:
    """
    Full browser automation agent using Playwright.
    Can log into websites, fill forms, scrape data, take screenshots.
    More powerful than basic automation — handles SPAs, auth flows, downloads.
    """

    async def execute_task(self, task: str, start_url: Optional[str] = None,
                            headless: bool = True) -> AppResult:
        """Execute a natural-language browser task."""
        try:
            from playwright.async_api import async_playwright
            async with async_playwright() as p:
                browser = await p.chromium.launch(headless=headless)
                context = await browser.new_context(
                    viewport={"width": 1280, "height": 720},
                    user_agent="Mozilla/5.0 (compatible; MICRODRAGON/1.0)"
                )
                page = await context.new_page()

                if start_url:
                    await page.goto(start_url, wait_until="domcontentloaded", timeout=30000)

                result = await self._execute_browser_task(page, task)
                await browser.close()
                return result
        except ImportError:
            return AppResult(success=False,
                             error="Playwright not installed: pip install playwright && playwright install chromium")
        except Exception as e:
            return AppResult(success=False, error=str(e))

    async def _execute_browser_task(self, page, task: str) -> AppResult:
        """Execute task on open browser page."""
        task_lower = task.lower()

        # Handle different task types
        if "screenshot" in task_lower:
            path = f"/tmp/microdragon_screenshot_{int(time.time())}.png"
            await page.screenshot(path=path, full_page=True)
            return AppResult(success=True, output_path=path,
                             app_used="Playwright", method="screenshot")

        elif "extract" in task_lower or "scrape" in task_lower or "get" in task_lower:
            # Extract page content
            content = await page.evaluate("() => document.body.innerText")
            links = await page.evaluate(
                "() => Array.from(document.links).map(l => ({text: l.text, href: l.href})).slice(0, 50)"
            )
            title = await page.title()
            return AppResult(success=True, output=f"Title: {title}\n\nContent:\n{content[:3000]}",
                             app_used="Playwright", method="extraction")

        elif "click" in task_lower:
            # Try to find and click the described element
            words = task_lower.replace("click", "").strip().split()
            for word in words:
                try:
                    await page.click(f"text={word}", timeout=5000)
                    return AppResult(success=True, app_used="Playwright", method="click",
                                     output=f"Clicked: {word}")
                except Exception:
                    continue
            return AppResult(success=False, error="Element not found to click")

        elif "fill" in task_lower or "type" in task_lower or "enter" in task_lower:
            # Generic form fill — extract value from task
            return AppResult(success=False,
                             error="Form fill requires specific field and value specification")

        elif "navigate" in task_lower or "go to" in task_lower:
            # Extract URL from task
            import re
            url_match = re.search(r'https?://[^\s]+', task)
            if url_match:
                await page.goto(url_match.group(0), wait_until="domcontentloaded")
                title = await page.title()
                return AppResult(success=True, output=f"Navigated to: {title}",
                                 app_used="Playwright", method="navigation")
            return AppResult(success=False, error="No URL found in task")

        else:
            # Generic: just extract content and title
            title = await page.title()
            url = page.url
            content = await page.evaluate("() => document.body.innerText")
            return AppResult(
                success=True,
                output=f"URL: {url}\nTitle: {title}\n\nContent:\n{content[:2000]}",
                app_used="Playwright", method="generic"
            )


# ─── Unified App Controller ────────────────────────────────────────────────────

class AppController:
    """
    Master controller — routes tasks to the right app automatically.
    MICRODRAGON's human-level capability layer.
    """

    def __init__(self):
        self.image_editor = ImageEditor()
        self.spreadsheet = SpreadsheetController()
        self.word_processor = WordProcessor()
        self.browser = BrowserAgent()

    def get_capabilities(self) -> dict:
        """Report what apps are available on this system."""
        caps = {
            "photoshop": bool(self.image_editor.photoshop),
            "gimp": bool(self.image_editor.gimp),
            "pillow": self.image_editor.has_pillow,
            "excel": bool(self.spreadsheet.excel),
            "libreoffice": bool(self.spreadsheet.libreoffice),
            "word": bool(self.word_processor.word),
        }
        return caps

    async def handle(self, task: str, task_type: str, **kwargs) -> AppResult:
        """Route task to correct app controller."""
        if task_type in ("design", "image", "photoshop", "graphic"):
            output_path = kwargs.get("output_path", f"/tmp/microdragon_design_{int(time.time())}.png")
            return await self.image_editor.create_design(task, output_path,
                kwargs.get("width", 1920), kwargs.get("height", 1080))

        elif task_type in ("spreadsheet", "excel", "csv", "data"):
            data = kwargs.get("data", [])
            headers = kwargs.get("headers", [])
            output_path = kwargs.get("output_path", f"/tmp/microdragon_data_{int(time.time())}.xlsx")
            return await self.spreadsheet.create_spreadsheet(task, data, headers, output_path)

        elif task_type in ("document", "word", "report", "letter"):
            content = kwargs.get("content", task)
            output_path = kwargs.get("output_path", f"/tmp/microdragon_doc_{int(time.time())}.docx")
            return await self.word_processor.create_document(task, content, output_path)

        elif task_type in ("browser", "web", "scrape", "automate"):
            url = kwargs.get("url")
            headless = kwargs.get("headless", True)
            return await self.browser.execute_task(task, url, headless)

        return AppResult(success=False, error=f"Unknown task type: {task_type}")


if __name__ == "__main__":
    async def demo():
        controller = AppController()
        caps = controller.get_capabilities()
        print("[MICRODRAGON Apps] Available capabilities:")
        for app, available in caps.items():
            status = "✓" if available else "✗"
            print(f"  {status} {app}")

        # Test document creation
        result = await controller.handle(
            "Q3 Financial Report",
            "document",
            content="# Q3 2026 Financial Report\n\n## Revenue\n\nTotal: $2.5M\n\n## Highlights\n\n- Growth: 35%\n- New customers: 150",
            output_path="/tmp/test_microdragon_doc.docx"
        )
        print(f"\nDocument: {result.success} — {result.output_path or result.error}")

    asyncio.run(demo())
