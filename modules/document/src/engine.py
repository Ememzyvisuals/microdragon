"""
microdragon/modules/document/src/engine.py
MICRODRAGON Document Intelligence — Read, analyze, and query any document
PDFs, DOCX, XLSX, presentations, text files — all queryable via AI
"""

import asyncio
import os
import tempfile
from dataclasses import dataclass
from typing import Optional
from pathlib import Path


@dataclass
class DocumentResult:
    success: bool
    content: str = ""
    page_count: int = 0
    word_count: int = 0
    file_type: str = ""
    summary: Optional[str] = None
    error: str = ""


class PDFReader:
    """Extract text and structure from PDFs."""

    def extract(self, pdf_path: str) -> DocumentResult:
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(pdf_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text.strip())

            content = "\n\n".join(pages)
            return DocumentResult(
                success=True, content=content,
                page_count=page_count,
                word_count=len(content.split()),
                file_type="pdf"
            )
        except ImportError:
            return self._extract_pypdf2(pdf_path)
        except Exception as e:
            return DocumentResult(success=False, error=str(e))

    def _extract_pypdf2(self, pdf_path: str) -> DocumentResult:
        try:
            import PyPDF2
            with open(pdf_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages = [page.extract_text() or "" for page in reader.pages]
            content = "\n\n".join(p.strip() for p in pages if p.strip())
            return DocumentResult(success=True, content=content,
                                  page_count=len(pages), word_count=len(content.split()),
                                  file_type="pdf")
        except ImportError:
            # Last resort: pdftotext CLI
            return self._extract_cli(pdf_path)
        except Exception as e:
            return DocumentResult(success=False, error=str(e))

    def _extract_cli(self, pdf_path: str) -> DocumentResult:
        import subprocess
        try:
            tmp = tempfile.mktemp(suffix=".txt")
            subprocess.run(["pdftotext", pdf_path, tmp], capture_output=True, timeout=30)
            if os.path.exists(tmp):
                content = Path(tmp).read_text(encoding="utf-8", errors="replace")
                os.unlink(tmp)
                return DocumentResult(success=True, content=content,
                                      word_count=len(content.split()), file_type="pdf")
        except Exception:
            pass
        return DocumentResult(success=False, error="No PDF reader available. pip install pdfplumber")


class DocxReader:
    """Read Word documents."""

    def extract(self, docx_path: str) -> DocumentResult:
        try:
            from docx import Document
            doc = Document(docx_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cell_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if cell_text:
                        paragraphs.append(cell_text)

            content = "\n\n".join(paragraphs)
            return DocumentResult(success=True, content=content,
                                  word_count=len(content.split()), file_type="docx",
                                  page_count=len(doc.paragraphs))
        except ImportError:
            return DocumentResult(success=False, error="python-docx not installed: pip install python-docx")
        except Exception as e:
            return DocumentResult(success=False, error=str(e))


class ExcelReader:
    """Read spreadsheets as structured text."""

    def extract(self, file_path: str) -> DocumentResult:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True, data_only=True)
            sheets = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(max_row=1000, values_only=True):
                    if any(cell is not None for cell in row):
                        rows.append(" | ".join(str(c) if c is not None else "" for c in row))
                if rows:
                    sheets.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))

            content = "\n\n".join(sheets)
            return DocumentResult(success=True, content=content,
                                  word_count=len(content.split()),
                                  file_type="xlsx", page_count=len(wb.sheetnames))
        except ImportError:
            return self._extract_csv_fallback(file_path)
        except Exception as e:
            return DocumentResult(success=False, error=str(e))

    def _extract_csv_fallback(self, path: str) -> DocumentResult:
        import csv
        try:
            rows = []
            with open(path, newline="", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    if i > 1000: break
                    rows.append(" | ".join(row))
            content = "\n".join(rows)
            return DocumentResult(success=True, content=content,
                                  word_count=len(content.split()), file_type="csv")
        except Exception as e:
            return DocumentResult(success=False, error=str(e))


class DocumentEngine:
    """Unified document reader and AI query engine."""

    def __init__(self):
        self.pdf = PDFReader()
        self.docx = DocxReader()
        self.xlsx = ExcelReader()

    def read(self, file_path: str) -> DocumentResult:
        """Auto-detect file type and extract content."""
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return self.pdf.extract(file_path)
        elif ext in (".docx", ".doc"):
            return self.docx.extract(file_path)
        elif ext in (".xlsx", ".xls", ".csv"):
            return self.xlsx.extract(file_path)
        elif ext in (".txt", ".md", ".rst", ".log", ".json", ".yaml", ".yml"):
            try:
                content = Path(file_path).read_text(encoding="utf-8", errors="replace")
                return DocumentResult(success=True, content=content,
                                      word_count=len(content.split()), file_type=ext[1:])
            except Exception as e:
                return DocumentResult(success=False, error=str(e))
        else:
            return DocumentResult(success=False, error=f"Unsupported file type: {ext}")

    def build_query_prompt(self, file_path: str, question: str, doc: DocumentResult) -> str:
        """Build a prompt for AI to answer questions about a document."""
        # Truncate content to fit context window
        content_limit = 15000
        content = doc.content[:content_limit]
        truncated = len(doc.content) > content_limit

        return f"""You are analyzing a document to answer a question.

**Document:** {Path(file_path).name}
**Type:** {doc.file_type.upper()}
**Size:** {doc.word_count:,} words, {doc.page_count} pages
{f"*Note: Document truncated to first {content_limit} chars*" if truncated else ""}

**Document Content:**
---
{content}
---

**Question:** {question}

Please answer the question based on the document content above. 
Be specific and cite relevant sections. If the answer is not in the document, say so clearly."""

    def summarize_prompt(self, file_path: str, doc: DocumentResult) -> str:
        """Build a summarization prompt."""
        content = doc.content[:20000]
        return f"""Summarize this document concisely:

**File:** {Path(file_path).name} ({doc.file_type}, {doc.word_count:,} words)

**Content:**
{content}

Provide:
1. Executive summary (2-3 sentences)
2. Key points (5-10 bullet points)
3. Important data/figures mentioned
4. Conclusions or recommendations"""


if __name__ == "__main__":
    engine = DocumentEngine()
    print("[MICRODRAGON Document] Engine ready")
    print("  Supported: PDF, DOCX, XLSX, CSV, TXT, MD, JSON, YAML")

    # Quick test
    import sys
    if len(sys.argv) > 1:
        result = engine.read(sys.argv[1])
        if result.success:
            print(f"  Read: {result.word_count} words, {result.page_count} pages")
            print(f"  Preview: {result.content[:200]}...")
        else:
            print(f"  Error: {result.error}")
